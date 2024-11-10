from src.utils.models import Parameters, SessionData
from docx import Document
from PyPDF2 import PdfReader
import re
import pandas as pd


class WordExtractor:
    def __init__(self, path: str, options: Parameters = Parameters(title=True)):
        self.options = options
        self._doc = Document(path)
        self._text = self.__get_text()
        self.file_data = {
            "title": self._extract_title,
            "technical_specification": self._extract_specification
        }

    def __get_text(self):
        data = ""
        for paragraph in self._doc.paragraphs:
            data += paragraph.text.lower() + "\n"
        return data

    def _extract_title(self):
        data = self._text
        title = data[data.find("поставк"):data.find("1.")].replace("\n", " ").strip()
        return title

    def _extract_specification(self):
        tables = []

        # Iterate through each table in the document
        for table in self._doc.tables:
            # Create a DataFrame structure with empty strings, sized by the number of rows and columns in the table
            df = [['' for _ in range(len(table.columns))] for _ in range(len(table.rows))]

            # Iterate through each row in the current table
            for i, row in enumerate(table.rows):
                # Iterate through each cell in the current row
                for j, cell in enumerate(row.cells):
                    # If the cell has text, store it in the corresponding DataFrame position
                    if cell.text:
                        df[i][j] = cell.text
            tables.append(pd.DataFrame(df))

        df = tables[0]
        slovar = {}
        for index, col in enumerate(df.T[0], start=0):
            if "характерис" in col.lower():
                slovar[index] = "Характеристики"
            elif "количество" in col.lower() or "кол-в" in col.lower():
                slovar[index] = "Количество"
            elif "наименов" in col.lower() or "назван" in col.lower():
                slovar[index] = "Название"
        df = df.rename(slovar, axis=1).drop(0)

        arr = []
        for row in df.itertuples():
            try:
                file_data = {
                    "Название": row.Название.lower(),
                    "Количество": re.search(r"\d+", str(row.Количество))[0],
                    "Характеристики": row.Характеристики.lower(),
                }
                arr.append(file_data)
            except Exception:
                return None
        return arr

    def parse(self):
        docx_data = SessionData()
        for option, value in self.options:
            if value:
                if not option in self.file_data:
                    continue
                setattr(docx_data, option, self.file_data[option]())
        return docx_data


class PdfExtractor:
    def __init__(self, path: str, options: Parameters = Parameters()):
        self.options = options
        self._doc = PdfReader(path)
        self._text = self.__get_text()
        self.file_data = {
            "title": self._extract_title,
        }

    def __get_text(self):
        data = ""
        for page in self._doc.pages:
            data += page.extract_text().lower() + "\n"
        return data

    def _extract_title(self):
        data = self._text
        return data[data.find("поставк"):data.find("г.")].replace("\n", " ").strip()

    def parse(self):
        pdf_data = SessionData()

        for option, value in self.options:
            if value:
                if not option in self.file_data:
                    continue
                setattr(pdf_data, option, self.file_data[option]())
        return pdf_data
